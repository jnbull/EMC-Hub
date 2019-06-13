"""To automate the generation of an EMC report based on project data

Recent software updates and bugfixes are documented here once software is in production

Software Version: a.001
Developer: Jadon Bull

"""

# ----------------------------Standard Imports-------------------------------
import os

# --------------------------Third-Party Imports------------------------------
import openpyxl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.shared import Cm
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_COLOR_INDEX
from copy import deepcopy

# ----------------------Global Variables & User Inputs---------------------
# these will eventually be django form inputs from the web app.
# productName1 = 'LPS-7'
# companyName1 = 'Virtek Vision'
# productClass1 = 'Class A'
# setup1 = 'Floor Standing'
# data1 = './EMC'
# temp1 = './static/reports/FCC.docx'
# output1 = './mainOutput.docx'

# ----------------------------Class Definitions------------------------------
class Report:

    """Used to create automated EMC reports based on test results.

    The "Report" class is a collector that contains instances of all tests
    conducted as attributes. Running the reportOutput function will create
    docx templates for all relevant tests.

    Attributes:
        product (str) - The name of the EUT this template belongs to
        company (str) - The name of the company this template belongs to
        class_ (str) - The class the EUT falls under. 'Class A' or 'Class B'
        testList (list) - A list of all tests completed for this report

        The following kwargs are optional:
        PLCE (class instance) - PLCE test data
        RE (class instance) - RE test data
        TLCE (class instance) - TLCE test data
        HARM (class instance) - Harmonics test data
        FLICK (class instance) - Flicker test data
        ESD (class instance) - ESD test data
        RI (class instance) - RI test data
        EFT (class instance) - EFT test data
        SURGE (class instance) - SURGE test data
        CI (class instance) - CI test data
        MI (class instance) - MI test data
        DIPS (class instance) - DIPS test data

    Public Methods:
        reportOutput() - fills docx templates with test data and outputs a
                         new completed docx file for each test inputted.


    """

    # --------------------------Initialization-------------------------------
    def __init__(self, productName, companyName, productClass, setup, data, temp, equipment, output, **kwargs):
        """Initializes a Report class instance for a particular project

        Parameters:
            productName (str) - The client-provided product name
            companyName (str) - The client's company name
            productClass (str) - The EUT's applicable class
            kwargs (dict) - Dictionary of template class instances

        """

        # Project Data
        self.product = productName
        self.company = companyName
        self.class_ = productClass
        self.setup = setup
        self.data = data
        self.temp = temp
        self.equipment = equipment
        self.output = output

        for key, value in kwargs.items():
            self.testList = []
            if key == 'PLCE':
                setattr(self, key, value(self.product, self.company, self.class_, self.setup, self.data, self.temp, self.equipment, self.output))
                self.testList.append(getattr(self, key))

    # --------------------------Public Methods-------------------------------
    def reportOutput(self):

        for test in self.testList:
            test.docOutput()

class Template:

    """Holds data relevant to a new EMC test report template.

    The 'Template' Class is used to organize important project data. When an
    EMC report template is generated, a Template instance will be created
    through inheritance of the corresponding test subclass. All template
    subclasses share the same attributes and methods that reside in this
    class.

    NOTE: The private methods within the class are meant to be used
    behind the scenes in order to generate a report template document
    using the python docx library.

    Attributes:
        product (str) - The name of the EUT this template belongs to
        company (str) - The name of the company this template belongs to
        class_ (str) - The class the EUT falls under. 'Class A' or 'Class B'

    """

    # --------------------------Initialization-------------------------------
    def __init__(self, productName, companyName, productClass, dataLocation):
        """Initializes a Template class instance with the project input data.

        Parameters:
            productName (str) - The client-provided product name
            companyName (str) - The client's company name
            productClass (str) - The EUT's applicable class
        """
        self.product = productName
        self.company = companyName
        self.class_ = productClass
        self.data = dataLocation

    # --------------------------Private Methods------------------------------
    def _getGIF(self, test):
        """Returns a list of GIF files from the project directory

        Output:
            fileList - a list of GIF file names
        """
        fileList = []

        dataDir = self.data + os.sep + 'Data'

        for dirName, subdirList, files in os.walk(dataDir):

            if test in subdirList:
                subdirList[:] = [test]
        
            breadCrumbs = dirName.split('/')

            for dir in breadCrumbs:
                if dir == 'DNU':
                    DNU = True
                else:
                    DNU = False

            if DNU == False:
                # print('Found Directory: %s' % dirName)
                for file in files:
                    filePath = dirName + os.sep + file
                    if filePath.endswith('.GIF'):
                        fileList.append(filePath)
        
        return(fileList)

    def _getXLSX(self, test):
        """Returns a list of XLSX files from the project directory.

        Output:
            fileList - a list of XLSX Detector file names
        """
        fileList = []

        dataDir = self.data + os.sep + 'Data'

        for dirName, subdirList, files in os.walk(dataDir):

            if test in subdirList:
                subdirList[:] = [test]
        
            breadCrumbs = dirName.split('/')

            for dir in breadCrumbs:
                if dir == 'DNU':
                    DNU = True
                else:
                    DNU = False

            if DNU == False:
                # print('Found Directory: %s' % dirName)
                for file in files:
                    filePath = dirName + os.sep + file
                    if filePath.endswith('.xlsx'):
                        fileList.append(filePath)
        
        return(fileList)

    def _removeEmptyRows(self, table):
        """Used to remove empty rows in table after data has been added.

        Parameters:
            table - the docx table object that rows will be removed from

        Output:
            (Implicit) rows are removed from table.

        """
        for row in table.rows:
            if row.cells[0].text == '0.000':
                self._remove_row(table, row)

    # --------------------------Static Methods-------------------------------
    @staticmethod
    def _copy_table_after(table, paragraph):
        tbl, p = table._tbl, paragraph._p
        new_tbl = deepcopy(tbl)
        p.addnext(new_tbl)

    @staticmethod
    def _insert_paragraph_before(item, text, style=None):
        p = CT_P.add_p_before(item._element)
        p2 = Paragraph(p, item._parent)
        p2.text = text
        p2.style = style
        return p2

    @staticmethod
    def _deleteParagraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    @staticmethod
    def _remove_row(table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)

    @staticmethod
    def _setRowHeight(table):
        for row in table.rows:
            row.height = Cm(0.53)

    @staticmethod
    def tableParse(detectorFile):
        header = []
        label = ''
        row = 1
        col = 1
        while True:
            if label == 'Angle':
                header.pop(-1)
                break
            else:
                label = detectorFile.cell(column=col, row=row).value
                col += 1
                header.append(label)

        # equipmentList(header)

        # If the first row entry has a detector labelled PEAK then only parse through 6 rows.
        peakList = dict()
        detectorType = detectorFile.cell(column=2, row=2).value
        detectorType = detectorType.split(' ')[0]
        print(detectorType)
        counter = 1
        if detectorType == 'SCAN':
            print('inside')
            rowRange = 8
            for row in range(2, rowRange):
                if not detectorFile.row_dimensions[row].hidden:
                    peakList[counter] = dict()
                    for col in range(1, len(header) + 1):
                        if col == 2:
                            label = 'PEAK'
                            peakList[counter][col] = label
                        elif col == 1:
                            peakList[counter][col] = format(detectorFile.cell(column=col, row=row).value, '.3f')
                        else:
                            peakList[counter][col] = format(detectorFile.cell(column=col, row=row).value, '.1f')
                    counter += 1
                else:
                    row += 1
                    rowRange += 1

        # If the first row entry has a detector other than PEAK (i.e. AVG or QP) then parse through all rows containing that label.
        else:
            print('outside')
            row = 2
            firstCellValue = 0
            counter = 1
            while firstCellValue != None:
                peakList[counter] = dict()
                for col in range(1, len(header) + 1):
                    cellValue = detectorFile.cell(column=col, row=row).value
                    if cellValue == None:
                        break
                    else:
                        if col == 1 or col == 2:
                            try:
                                peakList[counter][col] = format(cellValue, '.3f')
                            except ValueError:
                                peakList[counter][col] = cellValue
                        else:
                            peakList[counter][col] = format(cellValue, '.1f')
                row += 1
                counter += 1
                firstCellValue = detectorFile.cell(column=1, row=row).value

            # If there is less than 6, then get the remaining amount of PEAKS.
            detectorSum = len(peakList)
            print(detectorSum)
            maxSum = 6
            row = (detectorSum * 2) + 3
            while detectorSum < maxSum:
                print(row)
                if not detectorFile.row_dimensions[row].hidden:
                    peakList[counter] = dict()
                    for col in range(1, len(header) + 1):
                        if col == 2:
                            label = 'PEAK'
                            peakList[counter][col] = label
                        elif col == 1:
                            peakList[counter][col] = format(detectorFile.cell(column=col, row=row).value, '.3f')
                        else:
                            peakList[counter][col] = format(detectorFile.cell(column=col, row=row).value, '.1f')
                    detectorSum += 1
                    counter += 1
                else:
                    detectorSum += 1
                    maxSum += 1
                row += 1
        # print(detectorSum)
        # print(peakList)
        return (Template.marginHandler(peakList))

    @staticmethod
    def marginHandler(peakList):
        for row in range(1, len(peakList)):
            detectorType = peakList[row][2]
            if detectorType == 'AVG':
                peakList[row][8] = '--'
                peakList[row][10] = '--'
            elif detectorType == 'QP':
                peakList[row][9] = '--'
                peakList[row][11] = '--'
        return (peakList)

class PLCE(Template):

    """Builds PLCE test report template.

    The "PLCE" class is used to organize data corresponding to Conducted
    Emissions testing for a particular Project. The class also contains
    methods for generating a PLCE report template for use in a final
    report.

    NOTE: The private methods within the class are meant to be used behind
    the scenes in order to generate a report template document using the
    python docx library.

    Attributes:
        product (str) - Inherited from Template Class
        company (str) - Inherited from Template Class
        class_ (str) - Inherited from Template Class
        data (str) - File location for PLCE data
        doc (str) - File location for PLCE report template
        filesGIF (list) - List of GIF files in data
        filesXLSX (list) - List of XLSX files in data
        output (str) - File location for completed template output

    """

    def __init__(self, productName, companyName, productClass, setup, dataLocation, tempLocation, equipment, output):
        super().__init__(productName, companyName, productClass, dataLocation)
        self.setup = setup
        self.doc = Document(tempLocation)
        self.filesGIF = self.parseFiles(self._getGIF('PLCE'))
        self.filesXLSX = self.parseFiles(self._getXLSX('PLCE'))
        self.equipment = equipment
        self.output = output

    def classHandler(self):
        paragraphs = self.doc.paragraphs

        if self.class_ == 'Class A':
            for i in range(len(paragraphs)):
                runs = paragraphs[i].runs
                for run in runs:
                    # print(run.text)
                    if run.text == 'CLASS B':

                        self._deleteParagraph(paragraphs[i])
                        self._deleteParagraph(paragraphs[i + 1])
                        self._deleteParagraph(paragraphs[i + 2])
                        self._deleteParagraph(paragraphs[i + 3])
                        self.doc.tables[1]._element.getparent().remove(self.doc.tables[1]._element)

        if self.class_ == 'Class B':
            for i in range(len(paragraphs)):
                runs = paragraphs[i].runs
                for run in runs:
                    # print(run.text)
                    if run.text == 'CLASS A':

                        self._deleteParagraph(paragraphs[i])
                        self._deleteParagraph(paragraphs[i + 1])
                        self._deleteParagraph(paragraphs[i + 2])
                        self.doc.tables[0]._element.getparent().remove(self.doc.tables[0]._element)

    def docOutput(self):
        paragraphs = self.doc.paragraphs

        for paragraph in paragraphs:
            runs = paragraph.runs
            for run in runs:
                # print(run.text)
                if run.text == 'Insert PLCE Photos Here':
                    self.insertPhoto(paragraph)
                    self._deleteParagraph(paragraph)
        for i in range(len(self.sortDetectors())):
            for file in self.sortDetectors()[i]:
                print(file.name)
                wbPLCE = openpyxl.load_workbook(file.name, data_only=True)
                wsPLCE = wbPLCE.worksheets[0]
                self.tableName(self.insertTable(PLCE.tableParse(wsPLCE), file, i), file)
        for i in range(2, len(self.doc.tables) - 1):
            self._removeEmptyRows(self.doc.tables[i])
            PLCE._setRowHeight(self.doc.tables[i])
        self.classHandler()
        self.equipmentList()
        self.doc.save(self.output)
        

    def equipmentList(self):
        equipTable = self.doc.tables[-1]

        # Spectrum Analyzer Selection
        if self.equipment['SpecA'] == 'GEMC 160':
            for col in range(len(equipTable.rows[0].cells)):
                for paragraph in equipTable.rows[1].cells[col].paragraphs:
                    for run in paragraph.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        elif self.equipment['SpecA'] == 'GEMC 198':
            for col in range(len(equipTable.rows[0].cells)):
                for paragraph in equipTable.rows[2].cells[col].paragraphs:
                    for run in paragraph.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
        #LISN
        if self.equipment['LISN'] == 'GEMC 302':
            for col in range(len(equipTable.rows[0].cells)):
                for paragraph in equipTable.rows[3].cells[col].paragraphs:
                    for run in paragraph.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        elif self.equipment['LISN'] == 'GEMC 303':
            for col in range(len(equipTable.rows[0].cells)):
                for paragraph in equipTable.rows[4].cells[col].paragraphs:
                    for run in paragraph.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        elif self.equipment['LISN'] == 'GEMC 12X':
            for col in range(len(equipTable.rows[0].cells)):
                for paragraph in equipTable.rows[5].cells[col].paragraphs:
                    for run in paragraph.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        for i, row in enumerate(equipTable.rows):
            if i != 0:
                tableText = row.cells[0].paragraphs[0].runs[0].font
                # print(tableText.highlight_color)
                if tableText.highlight_color != WD_COLOR_INDEX.YELLOW:
                    self._remove_row(equipTable, row)
                else:
                    for col in range(len(equipTable.rows[0].cells)):
                        for paragraph in row.cells[col].paragraphs:
                            for run in paragraph.runs:
                                run.font.highlight_color = None

    def floorStanding(self):
        paragraphs = self.doc.paragraphs

        if self.setup == 'Floor Standing':
            # print('inside')
            for i in range(len(paragraphs)):
                runs = paragraphs[i].runs
                for run in runs:
                    # print(run.text)
                    if run.text == 'Typical Setup Diagram':
                        # print('inside')
                        p = paragraphs[i + 1]._insert_paragraph_before()
                        r = p.add_run()
                        r.add_break(WD_BREAK.LINE)
                        r.add_picture('./assets/images/FloorStanding.png', width=Cm(12.32), height=Cm(8.05))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        self._deleteParagraph(paragraphs[i + 2])
                        self._deleteParagraph(paragraphs[i + 3])

    def insertPhoto(self, paragraph):
        fileList = self.filesGIF
        nameList = self.sortPhotos()
        for name in nameList:
            for i in fileList:
                if fileList[i].name == name:
                    p = paragraph._insert_paragraph_before()
                    r = p.add_run()
                    # r.add_text(self.product)
                    # r.add_break(WD_BREAK.LINE)
                    if fileList[i].line == 'L1':
                        r.add_text('Line ' + '(' + fileList[i].line + ')' + ' ' + '-' + ' ' + fileList[i].voltage)
                        r.add_break(WD_BREAK.LINE)
                        r.add_break(WD_BREAK.LINE)
                        r.add_picture(name, width=Cm(16.51), height=Cm(12.38))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if name != nameList[-1]:
                            r.add_break(WD_BREAK.PAGE)

                    else:
                        r.add_text('Neutral ' + '(' + fileList[i].line + ')' + ' ' + '-' + ' ' + fileList[i].voltage)
                        r.add_break(WD_BREAK.LINE)
                        r.add_break(WD_BREAK.LINE)
                        r.add_picture(name, width=Cm(16.51), height=Cm(12.38))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if name != nameList[-1]:
                            r.add_break(WD_BREAK.PAGE)
        self.floorStanding()

    def insertTable(self, peakList, file, counter):
        self.tableGenerator()

        # value = file.value
        # valueString = str(value)
        # for char in value:
        #     print(char)
        #     if char == '0':
        #         singleDigit = True
        #         break
        #     else:
        #         singleDigit = False
        # print(singleDigit)
        # if singleDigit == True:
        #     valueSplit = valueString.split('0')
        #     num = int(valueSplit[1])
        # else:
        #     num = int(value)

        table = self.doc.tables[counter + 2]
        rowCounter = 1
        if file.line == 'L1':
            for row in range(5, len(table.rows)):
                colCounter = 0
                for col in range(len(table.rows[0].cells)):
                        # print(table.rows[row].cells[col].text)

                        try:
                            colCounter += 1
                            tableText = table.rows[row].cells[col]
                            tableText.text = str(peakList[rowCounter][colCounter])
                            tableText.paragraphs[0].runs[0].font.name = 'Calibri'
                            tableText.paragraphs[0].runs[0].font.size = Pt(11)
                            tableText.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except KeyError:
                            continue

                rowCounter += 1

        elif file.line == 'L2':
            for row in range(21, len(table.rows)):
                colCounter = 0
                for col in range(len(table.rows[0].cells)):
                    # print(table.rows[row].cells[col].text)

                    try:
                        colCounter += 1
                        tableText = table.rows[row].cells[col]
                        tableText.text = str(peakList[rowCounter][colCounter])
                        tableText.paragraphs[0].runs[0].font.name = 'Calibri'
                        tableText.paragraphs[0].runs[0].font.size = Pt(11)
                        tableText.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except KeyError:
                        continue

                rowCounter += 1

        return (table)

    def parseFiles(self, fileList):
        fileDict = {}
        key = 0

        for file in fileList:
            name = file
            parsed = file.split('/')
            parsed2 = parsed[-1].split('.')
            parsed3 = parsed2[0].split('_')

            value = parsed3[0]
            type = parsed3[1]
            voltage = parsed3[2]
            line = parsed3[-1]

            if line == 'L1-Detectors' or line == 'L2-Detectors':
                lineSplit = line.split('-')
                line = lineSplit[0]

            fileDict[key] = File(name, value, type, voltage, line)
            key += 1

        return fileDict

    def sortDetectors(self):
        pairList = []
        fileRemovalDict = {}

        fileList = self.filesXLSX
        for i in range(len(fileList)):
            try:
                file = self.filesXLSX[i]
                value = file.value
                # print(value)
                for char in value:
                    # print(char)
                    if char == '~' or char == '$':
                        valueValid = False
                        fileRemovalDict[i] = file
                        break
                    else:
                        valueValid = True

                if valueValid == True:
                    continue
            except KeyError:
                continue

        for key in fileRemovalDict:
            # print(key)
            fileList.pop(key)

        # print(fileList)

        for i in range(len(fileList) + 1):
            try:
                activeValue = fileList[i].value
                # print(activeValue)
                for j in range(len(fileList)):
                    if j == i:
                        continue
                    else:
                        comparedValue = fileList[j].value
                        if activeValue == comparedValue:
                            pair = [fileList[i], fileList[j]]
                            pairList.append(pair)
                            for pairs in pairList:
                                if pair == pairs[::-1]:
                                    pairList.pop(-1)
            except KeyError:
                continue

        # print(pairList)
        return (pairList)

    def sortPhotos(self):

        nameList = []
        for i in self.filesGIF:
            nameList.append(self.filesGIF[i].name)

        return(sorted(nameList))

    def tableGenerator(self):
        table = self.doc.tables[2]

        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if run.text == 'Insert Emissions Data Here':
                    for num in range(1, len(self.sortDetectors())):
                        PLCE._copy_table_after(table, paragraph)
                    for i in range(4, len(self.doc.tables) - 1):
                        p = PLCE._insert_paragraph_before(self.doc.tables[i], 'Average and Quasi-Peak Emissions Table')
                        r = p.runs[0]
                        r.add_break(WD_BREAK.PAGE)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.text = 'Average and Quasi-Peak Emissions Table'
                    run.add_break(WD_BREAK.PAGE)

    def tableName(self, table, file):

        row0 = table.rows[0]
        row0.cells[4].text = self.class_

        row1 = table.rows[1]
        row1.cells[4].text = self.product

        row2 = table.rows[2]
        row2.cells[4].text = file.voltage

        for i in range(3):
            table.rows[i].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.rows[i].cells[4].paragraphs[0].runs[0].font.name = 'Calibri'
            table.rows[i].cells[4].paragraphs[0].runs[0].font.size = Pt(11)

        row2.cells[4].paragraphs[0].runs[0].font.bold = True

class RE(Template):

    def __init__(self, productName, companyName, productClass, dataLocation, tempLocation):
        super().__init__(productName, companyName, productClass)
        self.data = dataLocation
        self.doc = Document(tempLocation)
        self.filesGIF = []
        self.filesXLSX = []

class File():

    def __init__(self, name, value, type, voltage, line, frequency = None, polarity = None):
        self.name = name
        self.value = value
        self.type = type
        self.voltage = voltage
        self.frequency = frequency
        self.line = line
        self.polarity = polarity



# def main():

#     report = Report(productName1, companyName1, productClass1, setup1, data1, temp1, output1, PLCE = PLCE)
#     report.reportOutput()

# main()