from openpyxl import load_workbook
from openpyxl.styles import Alignment
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def createForm(date, assetNumber, engineer, peakValue, riseTime, fallTime, burstPeriod, burstDuration, formLocation, output):
    doc = getTemplate(assetNumber, formLocation)
    infoTable = doc.tables[0]
    nameTable = doc.tables[1]
    EFTtable = doc.tables[3]
    peakValue = peakValue/1000

    infoTable.cell(0,1).text = 'Date: ' + str(date)

    if engineer == 'SV':
        nameTable.cell(0,0).text = 'Tester: Sanjiv Vyas'
    elif engineer == 'MM':
        nameTable.cell(0,0).text = 'Tester: Marty McLear'
    elif engineer == 'MX':
        nameTable.cell(0,0).text = 'Tester: Min Xie'
    elif engineer == 'JN':
        nameTable.cell(0,0).text = 'Tester: James Nash'
    elif engineer == 'AE':
        nameTable.cell(0,0).text = 'Tester: Amir Emami'
    elif engineer == 'RA':
        nameTable.cell(0,0).text = 'Tester: Ray Au'
    elif engineer == 'JB':
        nameTable.cell(0,0).text = 'Tester: Jadon Bull'

    EFTtable.cell(2,3).text = str(peakValue)
    EFTtable.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EFTtable.cell(3,3).text = str(riseTime)
    EFTtable.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EFTtable.cell(4,3).text = str(fallTime)
    EFTtable.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EFTtable.cell(5,3).text = str(burstPeriod)
    EFTtable.cell(5,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    EFTtable.cell(6,3).text = str(burstDuration)
    EFTtable.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER



    doc.save(output)

    
def getTemplate(assetNumber, formLocation):
    if assetNumber == 'GEMC 317':
        return Document(formLocation + 'Verification Form - GEMC 317.docx')

def addEFTRecord(date, assetNumber, engineer, peakValue, riseTime, fallTime, fileLocation):
    wb = load_workbook(fileLocation + 'EFT Verification Trends - V2.0.xlsx')
    ws = wb["Measurements"]
    # doc = Document('./static/files/Verification Form - GEMC 317')

    i = 4
    while ws.cell(row=i, column=1).value != None:
        i+=1
    print(i)

    dateCell = ws.cell(row=i, column=1)
    engineerCell = ws.cell(row=i, column=2)
    assetCell = ws.cell(row=i, column=3)
    appliedCell = ws.cell(row=i, column=4)
    peakCell = ws.cell(row=i, column=5)
    riseCell = ws.cell(row=i, column=6)
    fallCell = ws.cell(row=i, column=7)

    dateCell.value = date
    engineerCell.value = engineer
    assetCell.value = assetNumber
    appliedCell.value = 4000
    peakCell.value = peakValue
    riseCell.value = riseTime
    fallCell.value = fallTime

    dateCell.alignment = Alignment(horizontal='left')
    engineerCell.alignment = Alignment(horizontal='left')
    assetCell.alignment = Alignment(horizontal='left')
    appliedCell.alignment = Alignment(horizontal='left')
    peakCell.alignment = Alignment(horizontal='left')
    riseCell.alignment = Alignment(horizontal='left')
    fallCell.alignment = Alignment(horizontal='left')
    
    # print(ws.cell(row=5, column=1).number_format)

    wb.save(fileLocation + 'EFT Verification Trends - V2.0.xlsx')

def addESDRecord(date, assetNumber, posPeakValue, pos30Value, negPeakValue, neg30Value):
    wb = load_workbook('./static/files/ESD Verification Trends - V2.0.xlsx')
    ws = wb["Measurements"]

    i = 7
    while ws.cell(row=i, column=1).value != None:
        i+=1
    print(i)

    dateCell = ws.cell(row=i, column=1)
    assetCell = ws.cell(row=i, column=3)
    posPeakValueCell = ws.cell(row=i, column=5)
    pos30ValueCell = ws.cell(row=i, column=6)
    negPeakValueCell = ws.cell(row=i, column=8)
    neg30ValueCell = ws.cell(row=i, column=9)

    dateCell.value = date
    assetCell.value = assetNumber
    posPeakValueCell.value = posPeakValue
    pos30ValueCell.value = pos30Value
    negPeakValueCell.value = negPeakValue
    neg30ValueCell.value = neg30Value

    dateCell.alignment = Alignment(horizontal='left')
    assetCell.alignment = Alignment(horizontal='left')
    posPeakValueCell.alignment = Alignment(horizontal='left')
    pos30ValueCell.alignment = Alignment(horizontal='left')
    negPeakValueCell.alignment = Alignment(horizontal='left')
    neg30ValueCell.alignment = Alignment(horizontal='left')

    wb.save('./static/files/ESD Verification Trends - V2.0.xlsx')

    

# print(ws.cell(row=1,column=1).value)